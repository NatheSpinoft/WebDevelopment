from docx import Document

def ask_questions():
    questions = [
        "What is the problem?",
        "When did the problem occur?",
        "What steps reproduce the problem?",
        "What is the expected outcome?",
        "What is the actual outcome?",
        "What environment are you working in?",
        "What have you tried so far?",
        "Are there any related issues or dependencies?",
        "Have you consulted any documentation or resources?",
        "What is the urgency of the issue?",
        "What are the possible solutions or hypotheses?",
        "Who else might need to be involved?"
    ]
    
    answers = {}
    print("Please answer the following questions:")
    for question in questions:
        answers[question] = input(question + " ")
    
    return answers

def create_docx(answers, filename="troubleshooting_report.docx"):
    doc = Document()
    doc.add_heading('Troubleshooting Report', 0)
    
    for question, answer in answers.items():
        doc.add_heading(question, level=1)
        doc.add_paragraph(answer)
    
    doc.save(filename)
    print(f"Report saved as {filename}")

if __name__ == "__main__":
    answers = ask_questions()
    create_docx(answers)
